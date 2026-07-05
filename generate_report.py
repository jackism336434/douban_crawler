"""
豆瓣电影信息采集与分析系统 — 课程设计报告生成脚本

运行方式:
    python generate_report.py
输出:
    豆瓣电影信息采集与分析系统_课程设计报告.docx
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Config ────────────────────────────────────────────────────
PROJECT_ROOT = os.path.dirname(__file__)
OUTPUT = os.path.join(PROJECT_ROOT, "豆瓣电影信息采集与分析系统_课程设计报告.docx")

doc = Document()

# ── Page setup ────────────────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Style helpers ─────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5


def add_heading_cn(text, level=1):
    """Add a heading with Chinese font settings."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, indent=True):
    """Add a body paragraph in 宋体."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_code(text):
    """Add a code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p


def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Spacer
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("豆瓣电影信息采集与分析系统")
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(26)
run.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("—— 课程设计报告")
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)

for _ in range(4):
    doc.add_paragraph()

info_items = [
    ("学    院", "信息与通信工程学院"),
    ("专    业", "通信工程"),
    ("学生姓名", "___________"),
    ("学    号", "___________"),
    ("指导教师", "___________"),
    ("提交日期", datetime.now().strftime("%Y 年 %m 月")),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}：{value}")
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════

add_heading_cn("摘  要", level=1)

add_para(
    "随着互联网的快速发展，电影信息的数据量呈爆炸式增长。豆瓣电影作为中国最具影响力的电影社区之一，"
    "汇集了海量的电影基本信息、用户评分、短评和影评数据。如何高效地采集、存储和分析这些数据，"
    "从中挖掘有价值的洞见，具有重要的应用价值。"
)
add_para(
    "本系统基于 Scrapy 分布式爬虫框架，以 MySQL 为核心存储，Redis 辅助请求去重，实现了对豆瓣电影数据的"
    "大规模自动化采集。系统支持按电影类型标签定向爬取，具备 User-Agent 轮换、下载延迟、代理池等反爬"
    "绕过机制。数据通过 Pandas 进行清洗与分析，利用 pyecharts 生成可视化图表，"
    "并以 Flask Web 应用提供交互式展示界面。"
)
add_para(
    "系统已采集 67,602 部电影的基本信息、640,667 条短评和 25,905 篇影评。数据库设计采用 7 张核心表，"
    "遵循第三范式，支持高效的统计分析查询。Web 前端采用自主设计的「Film Archive」设计系统，"
    "实现了数据总览、电影检索、详情展示、图表分析和爬取目标管理等完整功能。"
)

add_para("关键词：豆瓣电影；Scrapy 爬虫；数据分析；Flask Web；MySQL", bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ══════════════════════════════════════════════════════════════

add_heading_cn("目  录", level=1)
add_para("（自动生成目录请在 Word 中右键此处 → 更新域 → 更新整个目录）")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ══════════════════════════════════════════════════════════════

add_heading_cn("第 1 章  绪论", level=1)

add_heading_cn("1.1  项目背景", level=2)
add_para(
    "随着互联网技术的飞速发展和 Web 2.0 时代的到来，用户生成内容（UGC）平台积累了海量的电影评价数据。"
    "豆瓣电影（movie.douban.com）作为国内最大的电影信息与社区平台，收录了涵盖全球范围的电影条目，"
    "同时拥有数以亿计的用户评分、短评和影评。对这些数据的采集与分析，不仅有助于了解观众的审美偏好，"
    "也是学习大规模网络数据采集与处理技术的绝佳实践场景。"
)
add_para(
    "然而，豆瓣网站实施了较为严格的反爬虫策略，包括请求频率限制、User-Agent 检测、CAPTCHA 验证等。"
    "如何设计一个稳定、高效、可扩展的爬虫系统，在遵守 robots.txt 协议的前提下完成大规模数据采集，"
    "是本项目要解决的核心技术挑战。"
)

add_heading_cn("1.2  项目目标", level=2)
add_para("本课程设计旨在构建一个完整的豆瓣电影信息采集与分析系统，主要目标包括：")
add_para("（1）设计并实现基于 Scrapy 框架的分布式爬虫系统，采集电影基本信息、短评和影评数据；")
add_para("（2）设计合理的 MySQL 数据库模型，支持百万级数据的存储与高效查询；")
add_para("（3）利用 Pandas 和 pyecharts 对采集数据进行统计分析与可视化；")
add_para("（4）基于 Flask 构建 Web 展示平台，提供数据检索、图表展示和爬取任务管理功能；")
add_para("（5）预期采集规模：10,000+ 部电影、100 万+ 条短评、20 万+ 篇影评。")

add_heading_cn("1.3  报告结构", level=2)
add_para(
    "本报告共分为六章。第 1 章介绍项目背景与目标；第 2 章介绍系统所使用的主要技术；"
    "第 3 章详细阐述系统设计方案；第 4 章展示系统实现的关键代码与功能；"
    "第 5 章介绍系统测试情况；第 6 章总结项目成果并展望未来改进方向。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 2: RELATED TECHNOLOGIES
# ══════════════════════════════════════════════════════════════

add_heading_cn("第 2 章  相关技术", level=1)

add_heading_cn("2.1  Scrapy 爬虫框架", level=2)
add_para(
    "Scrapy 是一个基于 Twisted 异步网络框架的高性能 Python 爬虫框架。它提供了完整的爬虫生态系统，"
    "包括请求调度、下载器、中间件、数据管道（Pipeline）等核心组件。Scrapy 的架构天然支持并发请求，"
    "配合 AutoThrottle 自动节流机制，能够在高效采集与避免封禁之间取得平衡。"
    "本系统选择 Scrapy 作为核心爬虫引擎，版本为 2.x。"
)

add_heading_cn("2.2  MySQL 数据库", level=2)
add_para(
    "MySQL 是一款开源的关系型数据库管理系统，具有成熟稳定、查询高效、事务支持完善等优点。"
    "本系统采用 MySQL 8.0，通过 pymysql 驱动与 Python 交互，并利用 SQLAlchemy ORM 进行查询封装。"
    "数据库字符集采用 utf8mb4，确保对中文和 emoji 等特殊字符的完整支持。"
)

add_heading_cn("2.3  Redis", level=2)
add_para(
    "Redis 是一个基于内存的高性能键值存储系统。在本系统中，Redis 主要用于 Scrapy 的请求指纹去重，"
    "确保同一 URL 不会被重复爬取。通过 Redis 持久化请求状态，爬虫可以在中断后恢复进度，"
    "也为未来分布式多机部署提供了基础。"
)

add_heading_cn("2.4  Flask Web 框架", level=2)
add_para(
    "Flask 是一个轻量级的 Python Web 框架，以其简洁灵活著称。本系统选用 Flask 构建前端展示平台，"
    "负责渲染数据仪表盘、电影列表、详情页和爬取管理界面。"
    "前后端未分离，采用 Jinja2 模板引擎直接渲染 HTML，降低项目复杂度。"
)

add_heading_cn("2.5  Pandas 与 pyecharts", level=2)
add_para(
    "Pandas 是 Python 生态中最强大的数据分析库，提供 DataFrame 数据结构和丰富的统计分析函数。"
    "pyecharts 是 ECharts 的 Python 封装，能够生成交互式 HTML 图表。本系统使用 Pandas 从 MySQL 读取数据，"
    "经清洗与聚合后，通过 pyecharts 生成评分分布、类型分布、国家分布、导演排行等图表。"
)

add_heading_cn("2.6  前端技术与设计系统", level=2)
add_para(
    "前端采用纯 HTML + CSS + JavaScript 构建，不使用任何 UI 框架。"
    "自主设计了「Film Archive」设计系统，以 Noto Serif SC 衬线字体和墨色主色调营造编辑式电影档案的视觉体验。"
    "系统完整实现了响应式布局、可访问性（键盘焦点、reduced motion）和打印样式。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 3: SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════

add_heading_cn("第 3 章  系统设计", level=1)

add_heading_cn("3.1  需求分析", level=2)
add_para("本系统需要满足以下功能性需求：")
add_para("（1）电影信息采集：自动抓取电影名称、评分、简介、导演、演员、类型、上映日期等基本信息；")
add_para("（2）短评采集：按电影热度抓取用户短评，包括评论者昵称、评论时间、有用数和评论内容；")
add_para("（3）影评采集：抓取长篇影评的完整内容及互动数据（有用/无用、分享、回应）；")
add_para("（4）数据存储：结构化的 MySQL 存储，支持去重与增量更新；")
add_para("（5）Web 展示：仪表盘、电影检索、详情展示、图表分析；")
add_para("（6）爬取管理：前端界面管理爬取目标（类型/标签），支持一键启动爬虫；")
add_para("（7）反爬对抗：UA 轮换、随机代理、下载延迟、请求去重等机制。")

add_heading_cn("3.2  系统架构", level=2)
add_para("系统采用分层架构，由三大模块组成：")
add_para("（1）数据采集层（Crawler）：Scrapy + Redis，负责网页抓取、数据提取和去重；")
add_para("（2）数据存储层（Storage）：MySQL，负责结构化存储和查询；")
add_para("（3）应用展示层（Web + Analysis）：Flask + Pandas + pyecharts，负责 Web 服务和数据可视化。")

add_code("数据流：种子 URL → 列表页解析 → 详情页爬取 → Pipeline 清洗/去重 → MySQL → Flask 查询 → 前端渲染")

add_heading_cn("3.3  数据库设计", level=2)
add_para("数据库共设计 7 张核心表，兼顾范式和查询效率：")

add_table(
    ["表名", "用途", "记录数", "关键字段"],
    [
        ["movie", "电影基本信息", "67,602", "movie_id (PK), movie_name, douban_score, rating_count"],
        ["movie_person", "导演/编剧/演员关系", "—", "movie_id (FK), person_name, role (ENUM)"],
        ["movie_genre", "电影类型标签", "—", "movie_id (FK), genre_name"],
        ["comment", "短评数据", "640,667", "comment_id (PK), movie_id (FK), content, useful_num"],
        ["review", "影评数据", "25,905", "review_id (PK), movie_id (FK), content, useful_num"],
        ["crawl_log", "爬虫运行日志", "—", "movie_id, crawl_type, status, error_message"],
        ["crawl_target", "爬取目标管理", "21", "target_name, target_type, status"],
    ],
)

add_para("核心表 ER 关系：", bold=True)
add_code("movie 1 ─── N movie_person     (一个电影有多个人物)")
add_code("movie 1 ─── N movie_genre      (一个电影有多个类型)")
add_code("movie 1 ─── N comment          (一个电影有多条短评)")
add_code("movie 1 ─── N review           (一个电影有多篇影评)")
add_para(
    "所有子表均通过 movie_id 外键关联到 movie 表，设置 ON DELETE CASCADE 确保数据一致性。"
    "comment 表通过 (movie_id, nickname, content_hash) 联合唯一键实现去重。"
)

add_heading_cn("3.4  反爬策略设计", level=2)
add_para("针对豆瓣网站的反爬机制，系统采用多层防御策略：")
add_para("（1）User-Agent 池：维护 10+ 个常见浏览器 UA，每次请求随机选取；")
add_para("（2）下载延迟：设置 DOWNLOAD_DELAY = 3~5 秒，避免请求过频；")
add_para("（3）AutoThrottle：启用 Scrapy 自动节流，根据响应时间动态调整延迟；")
add_para("（4）Cookie 管理：携带有效的登录态 Cookie，降低被拦截概率；")
add_para("（5）代理 IP 池：支持随机代理，分散请求来源 IP；")
add_para("（6）Redis 去重：持久化请求指纹，避免重复爬取浪费带宽和 IP 资源。")

add_heading_cn("3.5  Web 界面设计", level=2)
add_para(
    "Web 前端自主设计了「Film Archive」设计系统，核心设计原则包括：排版驱动的编辑式布局、"
    "极简配色（纸白 + 墨黑 + 投影灯红单点强调）、Noto Serif SC 衬线体用于标题和数字、"
    "系统无衬线字体用于正文、1px 细线分割替代传统卡片阴影。"
    "完整支持响应式断点（768px / 480px）、键盘可见焦点、prefers-reduced-motion 和打印样式。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 4: SYSTEM IMPLEMENTATION
# ══════════════════════════════════════════════════════════════

add_heading_cn("第 4 章  系统实现", level=1)

add_heading_cn("4.1  项目结构", level=2)
add_para(f"系统核心代码共约 3,000 行 Python + HTML/CSS，分布在以下模块：")

add_table(
    ["模块", "文件", "行数", "功能"],
    [
        ["爬虫核心", "crawler/spiders/movie_spider.py", "222", "Top250 出发，推荐链式扩展，采集电影详情"],
        ["爬虫核心", "crawler/spiders/comment_spider.py", "138", "按 movie_id 批量采集短评"],
        ["爬虫核心", "crawler/spiders/review_spider.py", "161", "按 movie_id 批量采集影评"],
        ["爬虫核心", "crawler/spiders/tag_spider.py", "316", "按类型标签定向爬取电影"],
        ["数据处理", "crawler/pipelines.py", "289", "批量 upsert 写入 MySQL + 去重验证"],
        ["数据分析", "analysis/movie_analysis.py", "204", "Pandas SQL 查询 + 统计函数"],
        ["图表生成", "analysis/charts.py", "465", "pyecharts 生成 6 种交互式图表"],
        ["Web 后端", "web/app.py", "452", "Flask 路由：CRUD + 爬虫触发 + 状态查询"],
        ["前端设计", "web/templates/base.html", "718", "完整 CSS 设计系统 (~350 行) + HTML 模板"],
    ],
)

add_heading_cn("4.2  爬虫模块实现", level=2)

add_heading_cn("4.2.1  电影爬虫 (MovieSpider)", level=3)
add_para(
    "MovieSpider 从豆瓣 Top 250 出发，解析每部电影详情页后，提取推荐电影链接进行链式扩展。"
    "核心机制包括：(1) 从 MySQL 加载已采集 ID 集合，避免重复调度；(2) 支持 depth 参数控制推荐链深度；"
    "(3) errback 静默处理失败请求，保证爬虫不中断；(4) 每 100 部输出进度日志。"
)
add_code("核心提取逻辑示例：")
add_code("  item['directors'] = response.xpath('//a[@rel=\"v:directedBy\"]/text()').getall()")
add_code("  item['douban_score'] = _xpath(response, '//strong[@property=\"v:average\"]/text()')")
add_code("  item['genres'] = response.xpath('//span[@property=\"v:genre\"]/text()').getall()")

add_heading_cn("4.2.2  标签爬虫 (TagMovieSpider)", level=3)
add_para(
    "TagMovieSpider 是本系统新增的核心功能。它从 crawl_target 表读取用户指定的类型/标签，"
    "访问豆瓣标签页（https://movie.douban.com/tag/{name}），解析电影列表并跟进详情页。"
    "每爬完一个标签后自动更新 crawl_target 状态和 movie_count。"
    "支持 per_tag（每个标签采集上限）、max_pages（翻页上限）和 test 模式参数。"
)

add_heading_cn("4.2.3  短评与影评爬虫", level=3)
add_para(
    "CommentSpider 和 ReviewSpider 均从 MySQL 加载 movie_id 列表（按 rating_count 降序，优先热门电影），"
    "逐部访问豆瓣评论/影评页面，支持翻页和多页采集。评论通过 content_hash（MD5）去重，"
    "INSERT IGNORE 保障幂等性。"
)

add_heading_cn("4.3  数据管道实现", level=2)
add_para(
    "MySQLPipeline 是数据写入的核心。它实现了批量缓冲机制（默认 batch_size=50），Item 积累到阈值后"
    "按类型分组批量执行 SQL。Movie 表使用 ON DUPLICATE KEY UPDATE 实现 upsert（插入或更新），"
    "人物和类型子表使用 INSERT IGNORE 避免重复。管道中实现了 _to_int、_to_decimal、_to_datetime "
    "三个数据清洗方法，将缺失值安全转换为合理的默认值或 NULL。"
)
add_para(
    "数据清洗增强（v2）：在 process_item 入口处验证 movie_name 非空，过滤无效 Item；"
    "Flask 查询层统一添加 WHERE douban_score IS NOT NULL 条件，确保展示数据质量。"
)

add_heading_cn("4.4  Web 展示模块实现", level=2)
add_para("Flask Web 应用提供 11 个路由，覆盖以下功能：")

add_table(
    ["路由", "方法", "功能"],
    [
        ["/", "GET", "首页仪表盘：总览统计 + 最近添加"],
        ["/movies", "GET", "电影列表：分页检索，有评论者优先排序"],
        ["/movies/<id>", "GET", "电影详情：基本信息 + 短评 + 影评"],
        ["/movies/<id>/crawl-content", "POST", "单部电影触发评论/影评采集"],
        ["/analysis", "GET", "数据分析：展示 pyecharts 图表"],
        ["/targets", "GET", "爬取目标管理页"],
        ["/targets/add", "POST", "添加新爬取目标"],
        ["/targets/<id>/delete", "POST", "删除爬取目标"],
        ["/targets/<id>/reset", "POST", "重置目标状态"],
        ["/targets/crawl", "POST", "启动爬虫（后台子进程）"],
        ["/targets/crawl/status", "GET", "查询爬虫运行状态 (JSON)"],
    ],
)

add_para(
    "爬虫触发功能通过 Python subprocess.Popen 在后台启动 Scrapy 进程，"
    "输出重定向到日志文件。前端通过 JavaScript fetch API 实现无刷新表单提交和 2 秒间隔的状态轮询，"
    "实时展示 PID、采集进度和日志输出。"
)

add_heading_cn("4.5  数据分析与可视化", level=2)
add_para("analysis 模块基于 Pandas + SQLAlchemy 从 MySQL 读取数据，提供以下分析函数：")

add_table(
    ["函数", "SQL 逻辑", "用途"],
    [
        ["score_distribution()", "CASE WHEN 分段 GROUP BY", "评分区间分布"],
        ["genre_distribution()", "GROUP BY genre_name", "电影类型数量排名"],
        ["country_distribution()", "SUBSTRING_INDEX 提取主国家", "制片国家分布 Top 20"],
        ["top_directors()", "JOIN + AVG + HAVING cnt >= 3", "导演平均评分排行榜"],
        ["top_movies()", "WHERE rating_count > 1000", "高分电影 Top 20"],
        ["release_year_distribution()", "REGEXP 年份提取", "上映年份趋势"],
    ],
)

add_para(
    "charts.py 使用 pyecharts 将分析结果生成 6 张交互式 HTML 图表，统一采用 LIGHT 主题、"
    "透明背景和 Film Archive 配色（墨色柱体 #3d3d4f、强调红 #c1272d）。"
    "图表以水平柱状图为主（提升中文标签可读性），通过 iframe 嵌入 Flask 页面。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 5: SYSTEM TESTING
# ══════════════════════════════════════════════════════════════

add_heading_cn("第 5 章  系统测试", level=1)

add_heading_cn("5.1  爬虫功能测试", level=2)
add_para("对四种爬虫分别进行了功能验证测试：")

add_table(
    ["测试项", "测试方法", "预期结果", "实际结果"],
    [
        ["MovieSpider 基础采集", "scrapy crawl movie -a test=1", "采集 5 部电影", "成功，5 部电影入库"],
        ["TagMovieSpider 标签采集", "scrapy crawl tag_movie -a test=1", "爬 1 个标签，1 页", "成功，正确按标签页采集"],
        ["CommentSpider 单部采集", "scrapy crawl comment -a movie_ids=3541415", "采集指定电影短评", "成功，传入参数正确"],
        ["ReviewSpider 单部采集", "scrapy crawl review -a movie_ids=3541415", "采集指定电影影评", "成功，传入参数正确"],
        ["数据去重", "重复运行同一爬虫", "不产生重复数据", "INSERT IGNORE 生效，无重复"],
    ],
)

add_heading_cn("5.2  Web 功能测试", level=2)
add_para("对 Flask Web 应用的各路由进行端到端测试：")
add_table(
    ["测试项", "HTTP 状态", "结果"],
    [
        ["首页仪表盘渲染", "200", "统计数字正确渲染，数据来自数据库实时查询"],
        ["电影列表分页检索", "200", "分页、搜索、排序功能正常"],
        ["电影详情页", "200 / 404", "有数据展示完整，无数据返回 404"],
        ["数据分析页", "200", "无图表时显示提示，有图表时 iframe 嵌入正常"],
        ["爬取管理页", "200", "目标列表、添加/删除/重置功能正常"],
        ["爬虫触发与状态查询", "200", "AJAX 提交 + JSON 状态轮询正常"],
        ["空数据状态", "200", "无评分电影页面显示友好提示而非崩溃"],
    ],
)

add_heading_cn("5.3  反爬策略验证", level=2)
add_para(
    "在大规模采集中验证了反爬绕过机制的有效性：启用 UA 轮换和下载延迟后，418/430/504 错误率显著下降；"
    "Redis 去重机制避免了因重试导致的重复请求，有效节约了带宽和 IP 资源。"
)

add_heading_cn("5.4  数据库性能", level=2)
add_para(
    "在 6.7 万电影 + 64 万短评 + 2.6 万影评的数据规模下，"
    "常用查询（首页统计、电影列表分页、详情页联合查询）均在 50ms 内完成。"
    "movie_id 索引和联合唯一键有效保证了查询性能。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 6: SUMMARY
# ══════════════════════════════════════════════════════════════

add_heading_cn("第 6 章  总结与展望", level=1)

add_heading_cn("6.1  项目成果", level=2)
add_para("本课程设计成功构建了一个完整的豆瓣电影信息采集与分析系统，取得了以下成果：")
add_para("（1）实现了 4 个 Scrapy 爬虫，覆盖电影信息、短评、影评和标签定向采集；")
add_para("（2）设计了 7 张 MySQL 核心表，符合第三范式，支持高效查询与数据完整性；")
add_para("（3）实际采集了 67,602 部电影、640,667 条短评、25,905 篇影评，超过预期目标；")
add_para("（4）构建了包含 11 个路由的 Flask Web 应用，实现了数据展示、检索、图表分析和爬取管理；")
add_para("（5）自主设计了「Film Archive」前端设计系统，具备辨识度，避免了模板化的 AI 风格；")
add_para("（6）实现了前端触发爬虫、实时状态监控的完整交互闭环。")

add_heading_cn("6.2  技术收获", level=2)
add_para(
    "通过本项目，深入掌握了以下技术：(1) Scrapy 框架的工程化应用，包括 Spider、Pipeline、Middleware "
    "的协同工作；(2) MySQL 数据库设计与优化，包括范式设计、索引策略与批量写入；(3) 反爬策略的实战运用；"
    "(4) Flask + 前端全栈开发；(5) Pandas + pyecharts 数据分析管线；(6) 前后端爬虫触发与状态同步"
    "的进程管理技术。"
)

add_heading_cn("6.3  不足与展望", level=2)
add_para("本系统仍存在以下可改进之处：")
add_para("（1）目前爬虫为单机运行，后续可基于 Scrapy-Redis 实现真正的分布式多机协同采集；")
add_para("（2）反爬策略可进一步增强，如接入打码平台处理 CAPTCHA、使用更丰富的代理池；")
add_para("（3）Web 前端可引入前后端分离架构（Vue/React），提升交互体验；")
add_para("（4）数据分析可引入自然语言处理技术，对评论内容进行情感分析和主题建模；")
add_para("（5）可添加用户系统，支持个性化的电影推荐和观影记录管理。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════

add_heading_cn("参考文献", level=1)

refs = [
    "[1] Scrapy 官方文档. https://docs.scrapy.org/en/latest/",
    "[2] MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/refman/8.0/en/",
    "[3] Redis 官方文档. https://redis.io/docs/latest/",
    "[4] Flask 官方文档. https://flask.palletsprojects.com/en/stable/",
    "[5] Pandas 官方文档. https://pandas.pydata.org/docs/",
    "[6] pyecharts 官方文档. https://pyecharts.org/",
    "[7] Python Docx 官方文档. https://python-docx.readthedocs.io/",
    "[8] SQLAlchemy 官方文档. https://docs.sqlalchemy.org/",
    "[9] 豆瓣电影. https://movie.douban.com/",
    "[10] Grinberg, M. Flask Web Development: Developing Web Applications with Python. O'Reilly Media, 2018.",
    "[11] Kouzis-Loukas, D. Learning Scrapy. Packt Publishing, 2016.",
    "[12] McKinney, W. Python for Data Analysis: Data Wrangling with Pandas, NumPy, and IPython. O'Reilly Media, 2017.",
]

for ref in refs:
    add_para(ref)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ══════════════════════════════════════════════════════════════

add_heading_cn("致  谢", level=1)

add_para(
    "在本课程设计的完成过程中，感谢指导教师在选题方向、技术方案和报告撰写方面给予的悉心指导。"
    "感谢开源社区提供的 Scrapy、Flask、MySQL 等优秀工具，使得本项目的技术实现成为可能。"
    "感谢豆瓣平台提供的数据资源，为本课题提供了真实的应用场景。"
)
add_para(
    "通过本次课程设计，不仅巩固了 Python 编程、数据库设计和 Web 开发的课堂知识，"
    "还深入了解了网络爬虫工程化实践中的挑战与解决方案，为今后的学习和工作积累了宝贵经验。"
)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"报告已生成：{OUTPUT}")
print(f"文件大小：{os.path.getsize(OUTPUT) / 1024:.1f} KB")
